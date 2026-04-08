# -*- coding: utf-8 -*-
"""
公共文档处理工具 - Word导入导出、文本统计、差异对比
"""

import os
import re
import difflib
import tempfile


class AuxTools:
    """公共文档处理能力集合"""

    def __init__(self, api_client=None):
        self.api = api_client

    _LATEX_SPECIAL_CHARS = {
        '\\': r'\textbackslash{}',
        '{': r'\{',
        '}': r'\}',
        '$': r'\$',
        '&': r'\&',
        '#': r'\#',
        '%': r'\%',
        '_': r'\_',
        '~': r'\textasciitilde{}',
        '^': r'\textasciicircum{}',
    }

    def _escape_latex(self, text: str) -> str:
        """转义 LaTeX 保留字符。"""
        return ''.join(self._LATEX_SPECIAL_CHARS.get(char, char) for char in str(text or ''))

    def _latex_heading_command(self, line: str) -> str:
        """将常见标题行映射为 LaTeX 节命令。"""
        stripped = str(line or '').strip()
        if not stripped:
            return ''

        markdown_match = re.match(r'^(#{1,3})\s+(.*)$', stripped)
        if markdown_match:
            level = len(markdown_match.group(1))
            title = self._escape_latex(markdown_match.group(2).strip())
            if level == 1:
                return rf'\section{{{title}}}'
            if level == 2:
                return rf'\subsection{{{title}}}'
            return rf'\subsubsection{{{title}}}'

        if re.match(r'^第[一二三四五六七八九十百千万\d]+[章节部分篇]\s*', stripped):
            return rf'\section{{{self._escape_latex(stripped)}}}'

        if re.match(r'^[一二三四五六七八九十百千万\d]+[、.．)]\s*', stripped):
            return rf'\subsection{{{self._escape_latex(stripped)}}}'

        if re.match(r'^\d+(?:\.\d+){0,2}\s+', stripped):
            level = stripped.count('.')
            command = 'section' if level == 0 else 'subsection' if level == 1 else 'subsubsection'
            return rf'\{command}{{{self._escape_latex(stripped)}}}'

        return ''

    def _latex_body_from_text(self, text: str) -> str:
        """将普通文本整理为 LaTeX 正文。"""
        lines = []
        paragraph_buffer = []

        def flush_paragraph():
            if not paragraph_buffer:
                return
            paragraph_text = '\n'.join(part.strip() for part in paragraph_buffer if part.strip())
            paragraph_buffer.clear()
            if paragraph_text:
                lines.append(self._escape_latex(paragraph_text))
                lines.append('')

        for raw_line in str(text or '').splitlines():
            stripped = raw_line.strip()
            if not stripped:
                flush_paragraph()
                continue

            heading_command = self._latex_heading_command(stripped)
            if heading_command:
                flush_paragraph()
                lines.append(heading_command)
                lines.append('')
                continue

            paragraph_buffer.append(raw_line.rstrip())

        flush_paragraph()
        return '\n'.join(lines).rstrip()

    def import_docx(self, filepath: str) -> str:
        """导入Word文档，返回文本内容"""
        try:
            normalized_path = os.path.abspath(os.path.expanduser(str(filepath or '').strip()))
            if not normalized_path.lower().endswith('.docx'):
                raise RuntimeError('仅支持导入 .docx 格式的 Word 文档')
            import docx
            doc = docx.Document(normalized_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            return '\n\n'.join(paragraphs)
        except ImportError:
            raise RuntimeError('请安装python-docx库: pip install python-docx')
        except Exception as e:
            raise RuntimeError(f'读取Word文件失败: {e}')

    def export_docx(self, text: str, filepath: str, title: str = '', level_font_styles: dict = None, sections_data: dict = None) -> bool:
        """导出为Word文档。sections_data 可包含 section_order, sections, section_levels 以结构化方式导出。"""
        try:
            import docx
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = docx.Document()

            # 设置页边距
            for section in doc.sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(3.17)
                section.right_margin = Cm(3.17)

            lfs = level_font_styles or {}
            body_style = lfs.get('body', {})
            h1_style = lfs.get('h1', {})
            h2_style = lfs.get('h2', {})
            h3_style = lfs.get('h3', {})

            body_font = body_style.get('font', '宋体')
            body_font_en = body_style.get('font_en', 'Times New Roman')
            body_pt = body_style.get('size_pt', 12)
            h1_font = h1_style.get('font', '黑体')
            h1_font_en = h1_style.get('font_en', 'Times New Roman')
            h1_pt = h1_style.get('size_pt', 16)
            h2_font = h2_style.get('font', '黑体')
            h2_font_en = h2_style.get('font_en', 'Times New Roman')
            h2_pt = h2_style.get('size_pt', 14)
            h3_font = h3_style.get('font', '黑体')
            h3_font_en = h3_style.get('font_en', 'Times New Roman')
            h3_pt = h3_style.get('size_pt', 12)

            from docx.oxml.ns import qn

            def _set_run_font(run, cn_font, en_font, pt_size):
                """设置 run 的中英文字体和字号"""
                run.font.name = en_font
                run.font.size = Pt(pt_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)

            # 设置默认字体
            style = doc.styles['Normal']
            style.font.name = body_font_en
            style.font.size = Pt(body_pt)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), body_font)

            # 添加标题
            if title:
                heading = doc.add_heading(title, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in heading.runs:
                    _set_run_font(run, h1_font, h1_font_en, h1_pt)

            # 结构化导出（有章节数据时）
            if sections_data and isinstance(sections_data, dict):
                sd_order = sections_data.get('section_order', [])
                sd_sections = sections_data.get('sections', {})
                sd_levels = sections_data.get('section_levels', {})
                if sd_order and sd_sections:
                    level_fonts = {
                        1: (h1_font, h1_font_en, h1_pt),
                        2: (h2_font, h2_font_en, h2_pt),
                        3: (h3_font, h3_font_en, h3_pt),
                    }
                    for sec_title in sd_order:
                        sec_body = str(sd_sections.get(sec_title, '') or '').strip()
                        sec_level = sd_levels.get(sec_title, 2)
                        # 写入章节标题
                        if sec_title.strip():
                            heading_level = min(max(sec_level, 1), 3)
                            p = doc.add_heading(sec_title.strip(), level=heading_level)
                            cn, en, pt = level_fonts.get(heading_level, (h2_font, h2_font_en, h2_pt))
                            for run in p.runs:
                                _set_run_font(run, cn, en, pt)
                        # 写入章节正文
                        if sec_body:
                            for para_text in sec_body.split('\n'):
                                if not para_text.strip():
                                    continue
                                p = doc.add_paragraph(para_text)
                                p.paragraph_format.first_line_indent = Pt(24)
                                p.paragraph_format.space_after = Pt(6)
                                for run in p.runs:
                                    _set_run_font(run, body_font, body_font_en, body_pt)
                    doc.save(filepath)
                    return True

            # 按段落添加内容（正则匹配模式，向后兼容）
            paragraphs = text.split('\n')
            for para_text in paragraphs:
                if not para_text.strip():
                    continue
                # 检测是否为一级标题行
                if re.match(r'^第[一二三四五六七八九十\d]+[章]', para_text):
                    p = doc.add_heading(para_text, level=1)
                    for run in p.runs:
                        _set_run_font(run, h1_font, h1_font_en, h1_pt)
                # 检测是否为二级标题行
                elif re.match(r'^[一二三四五六七八九十\d]+[、.．]', para_text) or \
                     re.match(r'^第[一二三四五六七八九十\d]+[节]', para_text) or \
                     re.match(r'^\d+\.\d+\s', para_text):
                    p = doc.add_heading(para_text, level=2)
                    for run in p.runs:
                        _set_run_font(run, h2_font, h2_font_en, h2_pt)
                # 检测是否为三级标题行
                elif re.match(r'^\d+\.\d+\.\d+\s', para_text) or \
                     re.match(r'^（[一二三四五六七八九十\d]+）', para_text):
                    p = doc.add_heading(para_text, level=3)
                    for run in p.runs:
                        _set_run_font(run, h3_font, h3_font_en, h3_pt)
                else:
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.first_line_indent = Pt(24)
                    p.paragraph_format.space_after = Pt(6)
                    for run in p.runs:
                        _set_run_font(run, body_font, body_font_en, body_pt)

            doc.save(filepath)
            return True
        except ImportError:
            raise RuntimeError('请安装python-docx库')
        except Exception as e:
            raise RuntimeError(f'导出Word失败: {e}')

    def export_doc(self, text: str, filepath: str, title: str = '', level_font_styles: dict = None, sections_data: dict = None) -> bool:
        """导出为 DOC 文档（依赖本机 Microsoft Word）"""
        tmp_docx = ''
        word = None
        document = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_docx = tmp_file.name

            self.export_docx(text, tmp_docx, title, level_font_styles=level_font_styles, sections_data=sections_data)

            try:
                import win32com.client  # type: ignore
            except Exception as exc:
                raise RuntimeError('导出 DOC 需要本机安装 Microsoft Word') from exc

            word = win32com.client.DispatchEx('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0
            document = word.Documents.Open(os.path.abspath(tmp_docx), ReadOnly=False, AddToRecentFiles=False)
            try:
                document.SaveAs2(os.path.abspath(filepath), FileFormat=0)
            except Exception:
                document.SaveAs(os.path.abspath(filepath), FileFormat=0)
            return True
        except RuntimeError:
            raise
        except Exception as e:
            raise RuntimeError(f'导出DOC失败: {e}')
        finally:
            if document is not None:
                try:
                    document.Close(False)
                except Exception:
                    pass
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            if tmp_docx and os.path.exists(tmp_docx):
                try:
                    os.remove(tmp_docx)
                except OSError:
                    pass

    def export_latex(self, text: str, filepath: str, title: str = '') -> bool:
        """导出为 LaTeX 源文件。"""
        try:
            safe_title = self._escape_latex(str(title or '').replace('\n', ' ').strip())
            body = self._latex_body_from_text(text)
            document_lines = [
                '% !TEX program = xelatex',
                r'\documentclass[UTF8]{ctexart}',
                r'\usepackage[a4paper,margin=2.54cm]{geometry}',
                r'\usepackage{setspace}',
                r'\setstretch{1.35}',
                '',
            ]

            if safe_title:
                document_lines.extend(
                    [
                        rf'\title{{{safe_title}}}',
                        r'\author{}',
                        r'\date{}',
                        '',
                    ]
                )

            document_lines.append(r'\begin{document}')
            if safe_title:
                document_lines.extend([r'\maketitle', ''])
            if body:
                document_lines.append(body)
            document_lines.extend(['', r'\end{document}', ''])

            with open(filepath, 'w', encoding='utf-8', newline='\n') as f:
                f.write('\n'.join(document_lines))
            return True
        except Exception as e:
            raise RuntimeError(f'导出LaTeX失败: {e}')

    def export_txt(self, text: str, filepath: str) -> bool:
        """导出为TXT文件"""
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(text)
            return True
        except Exception as e:
            raise RuntimeError(f'导出TXT失败: {e}')

    def export_pdf(self, text: str, filepath: str, title: str = '') -> bool:
        """导出为PDF（通过Word转换）"""
        # 先导出为docx，再尝试转PDF
        import tempfile
        tmp_docx = filepath.replace('.pdf', '_tmp.docx')
        try:
            self.export_docx(text, tmp_docx, title)
            # 尝试使用系统Word转换
            import subprocess
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'pdf', '--outdir',
                 os.path.dirname(filepath), tmp_docx],
                capture_output=True, timeout=30
            )
            if result.returncode == 0:
                return True
            raise RuntimeError('LibreOffice不可用，请手动从Word另存为PDF')
        except FileNotFoundError:
            raise RuntimeError('未找到PDF转换工具，请从Word文件另存为PDF')
        finally:
            if os.path.exists(tmp_docx):
                os.remove(tmp_docx)

    def diff_text(self, text1: str, text2: str) -> list:
        """对比两段文本，返回差异列表"""
        lines1 = text1.splitlines(keepends=True)
        lines2 = text2.splitlines(keepends=True)
        differ = difflib.unified_diff(
            lines1, lines2,
            fromfile='原文', tofile='修改后',
            lineterm=''
        )
        return list(differ)

    def diff_highlight(self, text1: str, text2: str) -> list:
        """返回带标记的差异列表，用于GUI高亮显示"""
        matcher = difflib.SequenceMatcher(None, text1, text2)
        result = []
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                result.append(('equal', text1[i1:i2]))
            elif tag == 'insert':
                result.append(('insert', text2[j1:j2]))
            elif tag == 'delete':
                result.append(('delete', text1[i1:i2]))
            elif tag == 'replace':
                result.append(('delete', text1[i1:i2]))
                result.append(('insert', text2[j1:j2]))
        return result

    def count_words(self, text: str) -> dict:
        """统计字数信息"""
        cn_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        en_words = len(re.findall(r'\b[a-zA-Z]+\b', text))
        numbers = len(re.findall(r'\b\d+\b', text))
        sentences = len(re.split(r'[。！？.!?]', text))
        paragraphs = len([p for p in text.split('\n') if p.strip()])
        return {
            'total': len(text),
            'chinese': cn_chars,
            'english_words': en_words,
            'numbers': numbers,
            'sentences': max(0, sentences - 1),
            'paragraphs': paragraphs,
        }

    def check_format(self, text: str, style: str = '学术论文') -> dict:
        """执行本地格式与结构检查。"""
        content = str(text or '')
        issues = []

        if ',' in content and '，' not in content:
            issues.append('建议使用中文逗号（，）替代英文逗号（,）')
        if '.' in content.replace('...', '') and '。' not in content:
            issues.append('建议使用中文句号（。）替代英文句号（.）')

        cn_nums = re.findall(r'[一二三四五六七八九十百千万]+', content)
        if cn_nums:
            issues.append(f'发现{len(cn_nums)}处中文数字，{style}建议优先使用阿拉伯数字')

        paragraphs = [p for p in content.split('\n') if p.strip()]
        short_paras = [p for p in paragraphs if len(p.strip()) < 50]
        if short_paras:
            issues.append(f'发现{len(short_paras)}个过短段落，建议合并或扩充')

        has_ref = bool(re.search(r'\[\d+\]', content))
        if len(content) > 500 and not has_ref:
            issues.append('未发现参考文献引用标注，建议补充引用编号')

        sentence_count = len([s for s in re.split(r'[。！？.!?]', content) if s.strip()])
        return {
            'issues': issues,
            'issue_count': len(issues),
            'word_count': len(content),
            'para_count': len(paragraphs),
            'sentence_count': sentence_count,
        }

    def detect_sensitive(self, text: str) -> list:
        """检测敏感内容"""
        sensitive_patterns = [
            (r'作弊|抄袭|代写|枪手', '学术诚信风险'),
            (r'政治|党|政府|领导人', '政治敏感内容'),
            (r'色情|暴力|恐怖', '违规内容'),
        ]
        found = []
        for pattern, category in sensitive_patterns:
            matches = re.findall(pattern, text)
            if matches:
                found.append({'category': category, 'matches': matches[:5]})
        return found
